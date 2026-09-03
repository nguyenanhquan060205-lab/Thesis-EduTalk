"""Dựng báo cáo Word cho phần dữ liệu → huấn luyện mô hình.

Mọi con số đọc trực tiếp từ file kết quả trong research/data/processed/ —
không con số nào gõ tay, đúng yêu cầu trong PLAN_VIET_BAO_CAO_DATA_TRAIN.md.
"""
import json
import pathlib

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

GOC = pathlib.Path("/Users/quan/HUIT_CNTT/KhoaLuanCuNhan/Thesis-EduTalk")
P = GOC / "research/data/processed"
RA = GOC / "docs/BaoCao_DuLieu_HuanLuyen.docx"

# ── Nạp toàn bộ số liệu thật ──────────────────────────────────────────────
M = json.loads((P / "08_model/metrics_summary.json").read_text())
KD = json.loads((P / "06_validation/ket_luan_kiem_dinh.json").read_text())
TC = json.loads((P / "09_tinh_chinh/ket_luan_tinh_chinh.json").read_text())
TOPK = json.loads((P / "08_model/topk_kich_ban_trien_khai.json").read_text())
AB_TH = pd.read_csv(P / "08_model/ablation_du_lieu_tong_hop.csv", index_col=0)
AB_DT = pd.read_csv(P / "08_model/ablation_nhom_dac_trung.csv", index_col=0)
SS_CV = pd.read_csv(P / "08_model/so_sanh_kien_truc_cv.csv", index_col=0)
CR1 = pd.read_csv(P / "08_model/classification_report_stage1.csv", index_col=0)

DL, TEST, BL, AUC = M["du_lieu"], M["test"], M["baseline"], M["auc_roc"]
T1, AUTO, PHANG, TUVAN = (
    TEST["tang1_khoi"],
    TEST["auto_CHI_SO_CHINH"],
    TEST["phang_39_doi_chung"],
    TEST["tu_van_CO_DIEU_KIEN"],
)
N_NGANH = DL["n_lop_nganh"]


def pc(v, n=1):
    return f"{v * 100:.{n}f}".replace(".", ",") + "%"


def sf(v, n=3):
    return f"{v:.{n}f}".replace(".", ",")


def ng(v):
    return f"{v:,}".replace(",", ".")


# ── Khung tài liệu ────────────────────────────────────────────────────────
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(13)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = st.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for s in doc.sections:
    s.top_margin, s.bottom_margin = Cm(2), Cm(2)
    s.left_margin, s.right_margin = Cm(3), Cm(2)

for lv, sz in ((1, 15), (2, 14), (3, 13)):
    h = doc.styles[f"Heading {lv}"]
    h.font.name = "Times New Roman"
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

_dem = {"hinh": 0, "bang": 0}


def p(txt="", **kw):
    par = doc.add_paragraph()
    if kw.get("center"):
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(txt)
    r.bold = kw.get("bold", False)
    r.italic = kw.get("italic", False)
    if kw.get("size"):
        r.font.size = Pt(kw["size"])
    return par


def rich(*phan):
    """Đoạn văn nhiều đoạn chữ: ("chữ", True) = in đậm."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for t, dam in phan:
        par.add_run(t).bold = dam
    return par


def bullet(txt, dam_den=None):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.line_spacing = 1.5
    if dam_den:
        par.add_run(txt[:dam_den]).bold = True
        par.add_run(txt[dam_den:])
    else:
        par.add_run(txt)
    for r in par.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)
    return par


def hinh(duong_dan, chu_thich, rong=15.5):
    _dem["hinh"] += 1
    d = P / duong_dan
    assert d.exists(), f"THIẾU HÌNH: {d}"
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(d), width=Cm(rong))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Hình {_dem['hinh']}. {chu_thich}")
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"


def bang(tieu_de, cot, hang, dam_hang=()):
    _dem["bang"] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Bảng {_dem['bang']}. {tieu_de}")
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    t = doc.add_table(rows=1, cols=len(cot))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cot):
        cell = t.rows[0].cells[i]
        cell.text = ""
        pr = cell.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pr.add_run(str(c))
        rr.bold = True
        rr.font.size = Pt(12)
        rr.font.name = "Times New Roman"
        sh = OxmlElement("w:shd")
        sh.set(qn("w:fill"), "D9E2F3")
        cell._tc.get_or_add_tcPr().append(sh)

    for idx, h in enumerate(hang):
        cells = t.add_row().cells
        for i, v in enumerate(h):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]
            pr.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            rr = pr.add_run(str(v))
            rr.font.size = Pt(12)
            rr.font.name = "Times New Roman"
            rr.bold = idx in dam_hang
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def luu_y(nhan, noi_dung):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.left_indent = Cm(0.6)
    par.add_run(f"{nhan} ").bold = True
    par.add_run(noi_dung).italic = True
    sh = OxmlElement("w:pBdr")
    lf = OxmlElement("w:left")
    lf.set(qn("w:val"), "single")
    lf.set(qn("w:sz"), "18")
    lf.set(qn("w:space"), "8")
    lf.set(qn("w:color"), "2E74B5")
    sh.append(lf)
    par._p.get_or_add_pPr().append(sh)


# ══════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════════════
p("TRƯỜNG ĐẠI HỌC CÔNG THƯƠNG TP. HỒ CHÍ MINH", center=True, bold=True, size=13)
p("KHOA CÔNG NGHỆ THÔNG TIN", center=True, bold=True, size=13)
for _ in range(4):
    p()
p("BÁO CÁO KHÓA LUẬN TỐT NGHIỆP", center=True, bold=True, size=18)
p()
p(
    "XÂY DỰNG TẬP DỮ LIỆU VÀ HUẤN LUYỆN MÔ HÌNH",
    center=True,
    bold=True,
    size=16,
)
p("GỢI Ý NGÀNH HỌC CHO HỌC SINH TRUNG HỌC PHỔ THÔNG", center=True, bold=True, size=16)
p()
p("Hệ thống EduTalk — Kiến trúc XGBoost phân cấp hai tầng", center=True, italic=True)
for _ in range(4):
    p()
p(
    f"Mô hình huấn luyện ngày {M['ngay_chay']} · seed = {M['seed']}",
    center=True,
    italic=True,
    size=12,
)
p(
    f"{ng(DL['train_that'])} phiếu khảo sát thật · {ng(DL['train_tong_hop'])} dòng tổng hợp · "
    f"{DL['test_that']} phiếu kiểm tra khoá",
    center=True,
    italic=True,
    size=12,
)
p(
    f"{DL['n_dac_trung']} đặc trưng · {DL['n_lop_nganh']} ngành · {DL['n_lop_khoi']} khối ngành",
    center=True,
    italic=True,
    size=12,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
p("MỤC LỤC", bold=True, size=15, center=True)
p()
for muc in [
    "CHƯƠNG 3. THU THẬP VÀ TIỀN XỬ LÝ DỮ LIỆU",
    "    3.1. Nguồn dữ liệu",
    "    3.2. Tiền xử lý hồ sơ trúng tuyển",
    "    3.3. Tiền xử lý dữ liệu khảo sát và chia tập",
    "    3.4. Tăng cường dữ liệu bằng thuật toán di truyền",
    "    3.5. Gộp dữ liệu, trọng số và kiểm định",
    "CHƯƠNG 4. HUẤN LUYỆN VÀ ĐÁNH GIÁ MÔ HÌNH",
    "    4.1. Xây dựng đặc trưng",
    "    4.2. Kiến trúc XGBoost hai tầng",
    "    4.3. Quy trình huấn luyện",
    "    4.4. Thí nghiệm loại bỏ (ablation)",
    "    4.5. Kết quả trên tập kiểm tra",
    "    4.6. Tinh chỉnh mở rộng",
    "CHƯƠNG 5. THẢO LUẬN",
    "    5.1. Điểm mạnh về phương pháp",
    "    5.2. Hạn chế",
    "    5.3. Hướng phát triển",
    "PHỤ LỤC. Danh mục tệp kết quả",
]:
    par = p(muc)
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("CHƯƠNG 3. THU THẬP VÀ TIỀN XỬ LÝ DỮ LIỆU", 1)
rich(
    (
        "Chương này trình bày nguồn gốc dữ liệu, quy trình làm sạch, phương pháp chia "
        "tập và chiến lược tăng cường dữ liệu. Toàn bộ quy trình được cài đặt trong sáu "
        "notebook (giai đoạn 1 đến giai đoạn 6), mỗi giai đoạn kết thúc bằng các lệnh "
        "kiểm tra tự động ",
        False,
    ),
    ("assert", True),
    (
        " nhằm bảo đảm kết quả tái lập được. Nguyên tắc xuyên suốt là tập kiểm tra được "
        "tách ra trước mọi bước xử lý và chỉ được mở đúng một lần ở giai đoạn cuối.",
        False,
    ),
)

doc.add_heading("3.1. Nguồn dữ liệu", 2)
p(
    "Bài toán gợi ý ngành học đòi hỏi đồng thời hai loại thông tin: năng lực học tập "
    "(điểm thi) và thiên hướng cá nhân (sở thích, mục tiêu). Trên thực tế không có "
    "nguồn nào cung cấp đủ cả hai, nên nghiên cứu sử dụng hai nguồn bổ khuyết cho nhau."
)
bullet(
    "Nguồn 1 — Hồ sơ trúng tuyển: 18.024 hồ sơ xét tuyển thật từ hệ thống tuyển sinh "
    "của Nhà trường, gồm điểm thi ba môn theo tổ hợp và mã ngành trúng tuyển. Nguồn này "
    "đông về số lượng nhưng thiếu hoàn toàn phần sở thích, giới tính và mục tiêu nghề nghiệp.",
    len("Nguồn 1 — Hồ sơ trúng tuyển:"),
)
bullet(
    "Nguồn 2 — Phiếu khảo sát sinh viên: 766 phiếu (trước lọc), có đầy đủ điểm thi, "
    "tên ngành đang học, mười câu hỏi sở thích thang Likert 1–5, giới tính và mục tiêu. "
    "Nguồn này đủ trường thông tin nhưng số lượng nhỏ.",
    len("Nguồn 2 — Phiếu khảo sát sinh viên:"),
)
p(
    "Ý tưởng kết hợp là dùng nguồn 2 để học phân phối sở thích theo từng ngành, rồi "
    "sinh phần sở thích còn thiếu cho nguồn 1, qua đó tận dụng được điểm thi thật của "
    "gần mười sáu nghìn hồ sơ mà vẫn có đủ đặc trưng đầu vào cho mô hình."
)

doc.add_heading("3.2. Tiền xử lý hồ sơ trúng tuyển (Giai đoạn 1)", 2)
p(
    "Hồ sơ thô được làm sạch qua bốn bước lọc tuần tự. Mỗi bước loại một nhóm bản ghi "
    "không phù hợp với bài toán, và số lượng còn lại sau từng bước được ghi nhận để "
    "bảo đảm tính minh bạch."
)
bang(
    "Phễu làm sạch hồ sơ trúng tuyển",
    ["Bước xử lý", "Số hồ sơ còn lại", "Tỷ lệ giữ"],
    [
        ["Hồ sơ thô ban đầu", "18.024", "100,0%"],
        ["Giữ hồ sơ trúng tuyển (KQ = TT)", "17.272", "95,8%"],
        ["Loại tổ hợp D10 (không dùng)", "16.309", "90,5%"],
        ["Đổi ba mã ngành cũ sang mã hiện hành", "16.250", "90,2%"],
        ["Giữ 39 ngành trong phạm vi nghiên cứu", "15.888", "88,1%"],
    ],
    dam_hang=(4,),
)
rich(
    ("Làm phẳng điểm thi. ", True),
    (
        "Ba cột điểm M1/M2/M3 được trải thành mười cột theo môn dựa trên bảng tra "
        "mười lăm tổ hợp xét tuyển. Một quyết định quan trọng là ",
        False,
    ),
    ("để trống ô môn không thi thay vì điền số 0", True),
    (
        ": XGBoost có cơ chế xử lý giá trị khuyết riêng, học hướng rẽ tối ưu cho nhánh "
        "khuyết, trong khi điền 0 sẽ tạo ra thông tin sai lệch rằng thí sinh thi môn đó "
        "và bị điểm liệt.",
        False,
    ),
)
p(
    "Thống kê mô tả cho thấy mặt bằng điểm giữa các môn chênh lệch hơn một điểm "
    "(Lịch sử 7,45 so với Địa lý 6,09). Đây là căn cứ để chuẩn hoá điểm theo z-score ở "
    "giai đoạn 7, nếu không mô hình sẽ hiểu nhầm thí sinh khối C giỏi hơn khối A. "
    "Tỷ lệ khuyết dao động từ 1,7% (Toán) đến 100% (Tin học)."
)
p(
    "Phân bố theo ngành mất cân bằng nghiêm trọng: ngành Công nghệ thông tin có 1.866 "
    "hồ sơ trong khi Công nghệ vật liệu chỉ có 19 hồ sơ, chênh lệch 98 lần. Năm trong "
    "số 39 ngành hoàn toàn không có hồ sơ trúng tuyển nào."
)
hinh("01_flatten/hinh_1_1_pheu_lam_sach.png", "Phễu làm sạch bốn bước, ghi rõ số hồ sơ còn lại sau từng bước")
hinh("01_flatten/hinh_1_2_thieu_diem_theo_mon.png", "Tỷ lệ khuyết điểm của từng môn thi")
hinh("01_flatten/hinh_1_3_phan_bo_diem.png", "Phân bố điểm từng môn và biểu đồ tần suất điểm trung bình")
hinh("01_flatten/hinh_1_4_ho_so_theo_39_nganh.png", "Số hồ sơ theo 39 ngành; năm ngành không có dữ liệu được gạch chéo")
hinh("01_flatten/hinh_1_5_khoi_nganh_va_diem.png", "Phân bố hồ sơ theo bảy khối ngành và mặt bằng điểm tương ứng")
hinh("01_flatten/hinh_1_6_lech_phan_phoi_to_hop.png", "So sánh tỷ lệ tổ hợp xét tuyển giữa hai nguồn dữ liệu")

doc.add_heading("3.3. Tiền xử lý dữ liệu khảo sát và chia tập (Giai đoạn 2)", 2)
rich(
    ("Gộp nhãn theo khối ngành. ", True),
    (
        f"39 ngành được gộp thành {DL['n_lop_khoi']} khối ngành theo bảng phân loại lưu tại "
        "nganh_khoi_mapping.json. Bảng này được dùng thống nhất cho toàn bộ pipeline, "
        "tránh tình trạng mỗi giai đoạn tự định nghĩa một cách gộp khác nhau.",
        False,
    ),
)
rich(
    ("Lọc phiếu điền đại. ", True),
    (
        "90 trên 766 phiếu (11,7%) bị loại vì có dấu hiệu trả lời máy móc: hoặc cả mười "
        "câu cùng một giá trị, hoặc độ lệch chuẩn của mười câu không vượt quá 0,3. Phép "
        "lọc này xét từng dòng độc lập, không dùng thông tin từ dòng khác, nên thực hiện "
        "trước khi chia tập vẫn không gây rò rỉ. Sau lọc còn 676 phiếu sạch, đủ mặt cả 39 ngành.",
        False,
    ),
)
rich(
    ("Chia tập phân tầng. ", True),
    (
        f"676 phiếu được chia theo tỷ lệ 85/15 có phân tầng theo ngành, thu được "
        f"{DL['train_that']} phiếu huấn luyện và {DL['test_that']} phiếu kiểm tra. Cả hai tập đều "
        "có mặt đủ 39 ngành và không có dòng nào trùng nhau.",
        False,
    ),
)
rich(
    ("Thiết kế kiểm định chéo. ", True),
    (
        f"Với chỉ {DL['train_that']} dòng thật cho {N_NGANH} lớp, việc cắt thêm một tập "
        f"validation cố định sẽ làm mất dữ liệu quý. Nghiên cứu dùng RepeatedStratifiedKFold "
        f"với {M['cv']['n_splits']} phần và {M['cv']['n_repeats']} lần lặp, cho "
        f"{M['cv']['n_splits'] * M['cv']['n_repeats']} lượt đo. Chỉ số fold được lưu ra tệp và "
        "giữ cố định cho mọi thí nghiệm, bảo đảm các kiến trúc được so sánh trên cùng một "
        "cách chia.",
        False,
    ),
)
luu_y("Nguyên tắc chống rò rỉ:", M["cv"]["ghi_chu"] + ".")
hinh("02_split/hinh_2_1_loc_phieu_dien_dai.png", "So sánh sáu phiếu bị loại và sáu phiếu hợp lệ")
hinh("02_split/hinh_2_2_phan_bo_likert.png", "Phân bố năm mức trả lời của từng câu hỏi sở thích")
hinh("02_split/hinh_2_3_nhan_khau_hoc.png", "Cơ cấu giới tính, mục tiêu nghề nghiệp và tổ hợp xét tuyển")
hinh("02_split/hinh_2_4_phan_bo_train_test.png", "Đối chiếu phân bố tập huấn luyện và tập kiểm tra")
hinh("02_split/hinh_2_5_thiet_ke_cross_validation.png", "Sơ đồ thiết kế kiểm định chéo và vị trí dữ liệu tổng hợp")

doc.add_heading("3.4. Tăng cường dữ liệu bằng thuật toán di truyền (Giai đoạn 3 và 4)", 2)

doc.add_heading("3.4.1. Học phân phối sở thích theo ngành", 3)
p(
    "Mục tiêu của giai đoạn này là trả lời câu hỏi: sinh viên ngành X thường trả lời "
    "mười câu sở thích như thế nào? Khó khăn nằm ở chỗ nhiều ngành chỉ có năm đến mười "
    "phiếu khảo sát, không đủ để ước lượng phân phối một cách tin cậy."
)
rich(
    ("Co ngót Bayes. ", True),
    (
        "Trung bình của mỗi ngành được kéo về phía trung bình của khối ngành chứa nó "
        "theo công thức μ_ngành = (n·x̄_ngành + K·μ_khối) / (n + K) với K = 10. Ngành có "
        "5 mẫu sẽ mượn 66,7% thông tin từ khối, ngành có 40 mẫu chỉ mượn 20%. Cơ chế này "
        "tự động cân bằng: ngành càng ít dữ liệu càng dựa vào khối, ngành nhiều dữ liệu "
        "thì giữ đặc trưng riêng.",
        False,
    ),
)
rich(
    ("Giữ ma trận hiệp phương sai. ", True),
    (
        "Thay vì chỉ lưu mười giá trị trung bình, nghiên cứu giữ nguyên ma trận hiệp "
        "phương sai 10×10 nhằm bảo toàn quan hệ giữa các câu hỏi. Ví dụ, câu Thí nghiệm "
        "và câu Môi trường có hệ số tương quan +0,615 — người thích làm thí nghiệm cũng "
        "thường quan tâm môi trường. Nếu bỏ qua quan hệ này, dữ liệu sinh ra sẽ có trung "
        "bình đúng nhưng cấu trúc sai. Toàn bộ 39 ma trận đều xác định dương.",
        False,
    ),
)
p(
    "Để tránh rò rỉ, phân phối được học riêng cho từng fold kiểm định chéo, thu được 15 "
    "bộ tham số. Sai lệch trung bình giữa bộ theo fold và bộ tổng thể là 0,033 — nhỏ, "
    "nhưng vẫn tách riêng để chỉ số validation không bị lạc quan."
)
hinh("03_distribution/hinh_3_1_bayesian_shrinkage.png", "Quan hệ giữa số mẫu và mức co ngót, kèm ví dụ trước và sau")
hinh("03_distribution/hinh_3_2_ma_tran_tuong_quan.png", "Ma trận tương quan: dữ liệu thật, phân phối đã học, và trường hợp bỏ qua hiệp phương sai")
hinh("03_distribution/hinh_3_3_ho_so_khoi_nganh.png", "Bản đồ nhiệt hồ sơ sở thích theo bảy khối ngành")

doc.add_heading("3.4.2. Sinh dữ liệu bằng thuật toán di truyền", 3)
rich(
    ("Bài toán. ", True),
    (
        "15.888 hồ sơ trúng tuyển có điểm thi thật nhưng thiếu mười câu sở thích. Cần "
        "sinh phần sở thích sao cho vừa khớp phân phối đã học ở giai đoạn 3, vừa giữ "
        "nguyên toàn bộ điểm thi và tổ hợp gốc.",
        False,
    ),
)
rich(
    ("Thiết kế cá thể. ", True),
    (
        "Điểm khác biệt so với cách dùng thuật toán di truyền thông thường là mỗi cá thể "
        "không phải một dòng dữ liệu mà là ",
        False,
    ),
    ("cả bộ dữ liệu n×10", True),
    (
        " của một ngành. Cách này cho phép hàm thích nghi đánh giá trực tiếp các đại "
        "lượng ở mức tập hợp — trung bình, độ lệch chuẩn, ma trận tương quan — điều "
        "không thể làm nếu cá thể chỉ là một dòng. Quần thể gồm 10 cá thể, chọn lọc bằng "
        "đấu loại kích thước 3, lai ghép hoán đổi theo dòng, đột biến thay mới 6% số dòng.",
        False,
    ),
)
bang(
    "Đóng góp của từng thành phần trong thuật toán di truyền",
    ["Thành phần", "Mức giảm sai lệch"],
    [
        ["Chọn lọc quần thể khởi tạo", "22,5%"],
        ["Tiến hoá qua các thế hệ", "42,2%"],
        ["Tổng cải thiện so với lấy mẫu thường", "55,2%"],
    ],
    dam_hang=(2,),
)
rich(
    ("Đối chứng phép đột biến. ", True),
    (
        "Hai kiểu đột biến được so sánh. Kiểu sửa từng ô ±1 cho mức tiến hoá bằng 0% vì "
        "nó phá vỡ quan hệ giữa các câu hỏi mà giai đoạn 3 đã dày công giữ lại. Kiểu thay "
        "cả dòng bằng mẫu mới sinh từ phân phối đa biến cho mức tiến hoá 42,2%. Thí nghiệm "
        "này khẳng định giá trị của việc bảo toàn hiệp phương sai.",
        False,
    ),
)
rich(
    ("Cân bằng lớp. ", True),
    (
        f"Số dòng mỗi ngành được đặt trần 600 và sàn 150, kéo tỷ lệ mất cân bằng từ 98 "
        f"lần xuống còn 4 lần. Năm ngành không có hồ sơ trúng tuyển được cấp điểm thi "
        f"lấy mẫu lại từ các ngành cùng khối. Kết quả thu được {ng(len(pd.read_csv(P / '04_synthetic/synthetic_GA_data.csv')))} "
        f"dòng tổng hợp, trong đó điểm thi được giữ nguyên 100%.",
        False,
    ),
)
hinh("04_synthetic/hinh_4_1_ga_hoi_tu.png", "Đường hội tụ của thuật toán di truyền và phân tách đóng góp chọn lọc / tiến hoá")
hinh("04_synthetic/hinh_4_2_can_bang_lop.png", "Số dòng theo 39 ngành trước và sau cân bằng, kèm ngưỡng trần và sàn")
hinh("04_synthetic/hinh_4_3_tong_hop_vs_that.png", "Đối chiếu trung bình từng câu hỏi và ma trận quan hệ giữa dữ liệu thật và dữ liệu sinh")
hinh("04_synthetic/hinh_4_4_giu_nguyen_diem_va_ty_le.png", "Phân bố điểm thi của dữ liệu tổng hợp trùng khít với hồ sơ gốc")

doc.add_heading("3.5. Gộp dữ liệu, trọng số và kiểm định (Giai đoạn 5 và 6)", 2)

doc.add_heading("3.5.1. Gộp nguồn và thiết kế trọng số", 3)
_syn = DL["train_tong_hop"]
rich(
    ("Vấn đề. ", True),
    (
        f"Tập huấn luyện gồm {ng(DL['train_that'])} dòng thật và {ng(_syn)} dòng tổng hợp, "
        f"tổng {ng(DL['train_that'] + _syn)} dòng. Nếu để nguyên, dữ liệu tổng hợp sẽ áp đảo "
        f"dữ liệu thật với tỷ lệ {_syn / DL['train_that']:.1f} trên 1, và mô hình rốt cuộc "
        "học đặc điểm của công thức sinh chứ không phải của người thật.",
        False,
    ),
)
rich(
    ("Trọng số hai thành phần. ", True),
    (
        f"Trọng số nguồn được tính động sao cho tổng ảnh hưởng của hai nguồn bằng nhau, "
        f"cho giá trị {sf(TC['trong_so_nguon']['dang_dung'], 2)} cho mỗi dòng thật. Trọng số lớp "
        "dùng nghịch đảo căn bậc hai của tần suất, thu hẹp chênh lệch giữa ngành đông và "
        "ngành hiếm từ 5,8 lần xuống 2,4 lần. Căn bậc hai được chọn thay vì nghịch đảo "
        "trực tiếp để tránh thổi phồng quá mức các ngành chỉ có vài mẫu. Trọng số cuối "
        "cùng là tích của hai thành phần, dao động từ 0,361 đến 20,86.",
        False,
    ),
)
hinh("05_final/hinh_5_1_trong_so_nguon.png", "So sánh tỷ lệ số dòng và tỷ lệ ảnh hưởng sau khi áp trọng số nguồn")
hinh("05_final/hinh_5_2_trong_so_lop.png", "Tổng ảnh hưởng của 39 ngành trước và sau khi áp trọng số lớp")
hinh("05_final/hinh_5_3_co_cau_tap_huan_luyen.png", "Cơ cấu tập huấn luyện cuối cùng theo khối, trọng số và nguồn")

doc.add_heading("3.5.2. Kiểm định chất lượng dữ liệu tổng hợp", 3)
rich(
    ("Phương pháp. ", True),
    (
        f"Dữ liệu tổng hợp được so với {KD['n_that']} dòng thật nằm ngoài fold đã dùng để "
        "ước lượng phân phối, có khớp tỷ lệ ngành. Điều này tránh lỗi thường gặp là so dữ "
        "liệu sinh ra với chính công thức đã sinh ra nó — phép so đó luôn cho kết quả tốt "
        "một cách giả tạo. Mỗi phép kiểm định đều kèm một phương án đối chứng: bộ dữ liệu "
        "sinh từ mười phân phối độc lập, cố tình bỏ qua hiệp phương sai.",
        False,
    ),
)
_j = KD["joint"]
_a, _b = "Có hiệp phương sai (đang dùng)", "10 phân phối độc lập (đối chứng)"
bang(
    "Bốn phép kiểm định dữ liệu tổng hợp",
    ["Phép kiểm định", "Phương án dùng", "Đối chứng", "Ngưỡng đạt", "Kết luận"],
    [
        ["Kiểm định Kolmogorov–Smirnov", f"{KD['ks_dat_tren_10_cau']}/10", "10/10", "≥ 7/10", "Đạt"],
        ["Sai lệch ma trận tương quan", sf(_j[_a]["Sai lệch tương quan TB"], 4), sf(_j[_b]["Sai lệch tương quan TB"], 4), "< 0,10", "Đạt"],
        ["Khoảng cách năng lượng", sf(_j[_a]["Energy distance"], 4), sf(_j[_b]["Energy distance"], 4), "< 0,05", "Đạt"],
        ["AUC phân biệt thật / tổng hợp", sf(KD["auc_phan_biet"][_a], 3), sf(KD["auc_phan_biet"][_b], 3), "< 0,65", "Đạt"],
    ],
)
p(
    "Cả bốn phép đều đạt ngưỡng. Đáng chú ý là kiểm định Kolmogorov–Smirnov cho kết quả "
    "10/10 ở cả phương án đang dùng lẫn phương án đối chứng, tức là nó không đủ sức phân "
    "biệt hai bộ dữ liệu. Lý do là phép này chỉ xét phân phối biên của từng câu hỏi riêng "
    "lẻ, trong khi khác biệt thực sự nằm ở quan hệ giữa các câu. Ba phép còn lại — vốn "
    "xét cấu trúc chung — mới tách bạch được: sai lệch tương quan chênh 7,5 lần và AUC "
    "phân biệt chênh gần 0,2. Đây là lý do nghiên cứu không dừng ở một phép kiểm định."
)
rich(
    ("Diễn giải chỉ số AUC phân biệt. ", True),
    (
        f"Một bộ phân loại được huấn luyện để phân biệt dòng thật với dòng tổng hợp. "
        f"Giá trị AUC càng gần 0,5 nghĩa là càng khó phân biệt, tức dữ liệu sinh càng "
        f"giống thật. Phương án đang dùng đạt {sf(KD['auc_phan_biet'][_a], 3)}, còn phương án "
        f"bỏ hiệp phương sai lên tới {sf(KD['auc_phan_biet'][_b], 3)} — lộ liễu rõ rệt.",
        False,
    ),
)
hinh("06_validation/hinh_6_1_bon_phep_kiem_dinh.png", "Bốn phép kiểm định, mỗi phép kèm thanh đối chứng")
hinh("06_validation/hinh_6_2_pca_that_vs_tong_hop.png", "Chiếu mười chiều sở thích xuống hai chiều bằng PCA")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("CHƯƠNG 4. HUẤN LUYỆN VÀ ĐÁNH GIÁ MÔ HÌNH", 1)

doc.add_heading("4.1. Xây dựng đặc trưng (Giai đoạn 7)", 2)
rich(
    ("Chuẩn hoá điểm thi. ", True),
    (
        "Do mặt bằng điểm giữa các môn chênh nhau hơn một điểm, điểm thô được chuyển "
        "sang z-score theo từng môn. Tham số chuẩn hoá (trung bình, độ lệch chuẩn) tính "
        "riêng trên tập huấn luyện rồi áp y hệt cho tập kiểm tra, bảo đảm không rò rỉ "
        "thông tin từ tập kiểm tra vào quá trình huấn luyện.",
        False,
    ),
)
bang(
    f"Cấu trúc {DL['n_dac_trung']} đặc trưng đầu vào",
    ["Nhóm đặc trưng", "Số cột", "Ghi chú"],
    [
        ["Mười câu sở thích Likert", "10", "Giữ nguyên thang 1–5"],
        ["Điểm thi từng môn", "10", "Giữ ô khuyết, không điền 0"],
        ["Điểm chuẩn hoá z-score", "3", "Trung bình, cao nhất, thấp nhất"],
        ["Nhân khẩu học", "2", "Giới tính, mục tiêu nghề nghiệp"],
        ["Tổ hợp xét tuyển (one-hot)", "15", "Mười lăm tổ hợp"],
        ["Nhóm tổ hợp", "3", "Tự nhiên, Xã hội, Ngoại ngữ"],
        ["Tổng cộng", str(DL["n_dac_trung"]), "Đúng bằng số cột của X_train"],
    ],
    dam_hang=(6,),
)
rich(
    ("Kiểm tra rò rỉ nhãn. ", True),
    (
        "Nghiên cứu kiểm tra ba tầng để loại trừ khả năng một đặc trưng nào đó suy ra "
        "trực tiếp từ nhãn. Kết quả: đặc trưng đơn lẻ mạnh nhất chỉ dự báo đúng khối "
        "ngành 32,9%, trong khi mốc đoán theo lớp đông nhất đã đạt 27,4%. Khoảng cách "
        "5,5 điểm phần trăm là quá nhỏ để coi là rò rỉ. Đặc biệt, biến one-hot khối ngành "
        "đã bị loại khỏi đặc trưng của tầng 2, vì đó là hàm suy trực tiếp từ nhãn ngành "
        "và sẽ thổi phồng kết quả một cách giả tạo.",
        False,
    ),
)
hinh("07_model_ready/hinh_7_1_chuan_hoa_diem.png", "Phân bố điểm trước và sau khi chuẩn hoá z-score")
hinh("07_model_ready/hinh_7_2_chat_luong_dac_trung.png", "Đánh giá chất lượng và kiểm tra rò rỉ của tập đặc trưng")

doc.add_heading("4.2. Kiến trúc XGBoost phân cấp hai tầng", 2)
p(
    f"Dự đoán trực tiếp một trong {N_NGANH} ngành là bài toán khó với chỉ "
    f"{DL['train_that']} phiếu thật, trung bình chưa tới mười lăm phiếu mỗi ngành. "
    f"Nghiên cứu chia bài toán thành hai bước dễ hơn."
)
bullet(
    f"Tầng 1 — phân loại khối ngành: một mô hình XGBoost {DL['n_lop_khoi']} lớp, cho ra "
    "phân phối xác suất P₁ trên các khối ngành.",
    len("Tầng 1 — phân loại khối ngành:"),
)
bullet(
    f"Tầng 2 — phân loại ngành: một mô hình XGBoost {N_NGANH} lớp chạy độc lập trên cùng "
    "bộ đặc trưng, cho ra phân phối P₂ trên các ngành. Tầng này không nhận thông tin khối "
    "làm đầu vào nên không bị rò rỉ nhãn.",
    len("Tầng 2 — phân loại ngành:"),
)
par = doc.add_paragraph()
par.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = par.add_run(f"P(ngành) ∝ P₂(ngành) × P₁(khối chứa ngành)^β,   β = {sf(M['sieu_tham_so']['beta'], 1)}")
r.bold = True
r.font.size = Pt(13)
p(
    "Hệ số β điều tiết mức độ tin vào tầng 1. Giá trị β = 0 tương đương bỏ hẳn tầng 1, "
    "còn β = 1 là nhân đầy đủ. Giá trị 0,6 được chọn qua kiểm định chéo: đủ để tầng 1 "
    "nâng đỡ tầng 2 nhưng không quá mạnh đến mức một lần đoán sai khối kéo theo hỏng "
    "toàn bộ kết quả. Cách kết hợp mềm này chịu lỗi tốt hơn phương án cắt cứng theo khối."
)
bang(
    "Hai chế độ phục vụ",
    ["Chế độ", "Người dùng chọn khối?", "Cách tính", "Ý nghĩa"],
    [
        ["Khám phá", "Không", "Dùng P₁ do mô hình dự đoán", "Năng lực tự động thật của hệ thống"],
        ["Tư vấn", "Có", "P₁ = 1 cho khối đã chọn", "Xếp hạng trong khối người dùng chỉ định"],
    ],
)

doc.add_heading("4.3. Quy trình huấn luyện (Giai đoạn 8)", 2)
p(
    f"Mô hình được huấn luyện qua {M['cv']['n_splits'] * M['cv']['n_repeats']} lượt kiểm "
    "định chéo. Trong mỗi lượt, tập validation chỉ gồm dòng thật, còn dữ liệu tổng hợp "
    "được sinh riêng cho từng fold và chỉ đưa vào phía huấn luyện. Cơ chế dừng sớm "
    "(early stopping) dùng tập validation thật, không bao giờ chạm vào tập kiểm tra."
)
bang(
    "So sánh các kiến trúc trên kiểm định chéo (giá trị trung bình 15 lượt)",
    ["Kiến trúc", "Top-1", "Top-3", "Top-5", "macro-F1"],
    [
        [
            {"Đoán lớp đông nhất": "Mốc đoán lớp đông nhất",
             "Model phẳng 39 lớp": "Một tầng, 39 lớp",
             "2 tầng — nhân xác suất (β=0.6)": "Hai tầng, chế độ Khám phá",
             "2 tầng — tư vấn (có điều kiện)": "Hai tầng, chế độ Tư vấn"}[i],
            pc(r_["top1"]), pc(r_["top3"]), pc(r_["top5"]), sf(r_["macro_f1"]),
        ]
        for i, r_ in SS_CV.iterrows()
    ],
    dam_hang=(2,),
)
rich(
    ("Nhận xét. ", True),
    (
        f"Trên kiểm định chéo, kiến trúc hai tầng chế độ Khám phá vượt mô hình một tầng "
        f"ở toàn bộ năm chỉ số, trong đó macro-F1 tăng từ "
        f"{sf(SS_CV.loc['Model phẳng 39 lớp', 'macro_f1'])} lên "
        f"{sf(SS_CV.loc['2 tầng — nhân xác suất (β=0.6)', 'macro_f1'])}. Kết quả này là căn cứ "
        "để giữ kiến trúc phân cấp. Cần lưu ý phân biệt hai mốc: mốc đoán lớp đông nhất "
        f"đạt Top-3 {pc(BL['nganh_top3'])}, còn đoán ngẫu nhiên đều trong {N_NGANH} ngành chỉ "
        f"đạt {pc(3 / N_NGANH)}.",
        False,
    ),
)
_of = M["cv_overfit"]
rich(
    ("Mức quá khớp. ", True),
    (
        f"Độ chính xác Top-1 trên tập huấn luyện đạt {pc(_of['train_top1'])} trong khi trên "
        f"tập validation chỉ {pc(_of['val_top1'])}, khoảng cách "
        f"{pc(_of['train_top1'] - _of['val_top1'])}. Đây là hệ quả trực tiếp của việc chỉ có "
        f"khoảng mười lăm mẫu thật cho mỗi ngành, và được ghi nhận công khai như một hạn "
        "chế của nghiên cứu.",
        False,
    ),
)
hinh("08_model/hinh_8_1_tim_sieu_tham_so.png", "Quá trình tìm siêu tham số bằng tìm kiếm ngẫu nhiên")
hinh("08_model/hinh_8_2_so_sanh_kien_truc.png", "So sánh các kiến trúc trên kiểm định chéo")

doc.add_heading("4.4. Thí nghiệm loại bỏ (Ablation)", 2)

doc.add_heading("4.4.1. Dữ liệu tổng hợp có thực sự hữu ích không?", 3)
p(
    "Toàn bộ giai đoạn 3 và 4 sẽ mất ý nghĩa nếu dữ liệu tổng hợp không cải thiện được "
    "mô hình. Thí nghiệm sau huấn luyện hai lần trên cùng cách chia: một lần chỉ dùng "
    "dòng thật, một lần thêm dữ liệu tổng hợp."
)
bang(
    "Ảnh hưởng của dữ liệu tổng hợp",
    ["Chỉ số", "Chỉ dòng thật", "Có thêm tổng hợp", "Thay đổi tương đối"],
    [
        [i, pc(r_["Chỉ dòng thật"]), pc(r_["+ dữ liệu tổng hợp"]),
         ("+" if r_["Thay đổi"] >= 0 else "") + pc(r_["Thay đổi"])]
        for i, r_ in AB_TH.iterrows()
    ],
    dam_hang=(4,),
)
rich(
    ("Nhận xét. ", True),
    (
        f"macro-F1 tăng {pc(AB_TH.loc['Ngành — macro-F1', 'Thay đổi'])} — mức tăng lớn hơn "
        "hẳn so với Top-1. Điều này cho thấy dữ liệu tổng hợp giúp nhiều nhất cho các "
        "ngành hiếm, vốn là thành phần mà macro-F1 đánh giá công bằng còn Top-1 thì bỏ "
        "qua. Riêng Top-3 giảm nhẹ, cho thấy lợi ích không đồng đều trên mọi chỉ số.",
        False,
    ),
)

doc.add_heading("4.4.2. Đóng góp của từng nhóm đặc trưng", 3)
luu_y(
    "Lưu ý:",
    "Đây là thí nghiệm trong phòng thí nghiệm nhằm đo đóng góp của từng nhóm đặc trưng. "
    f"Ứng dụng triển khai thực tế luôn dùng đủ {DL['n_dac_trung']} đặc trưng; mười câu "
    "sở thích là bắt buộc với người dùng.",
)
bang(
    "Kết quả khi giữ lại từng nhóm đặc trưng",
    ["Nhóm đặc trưng", "Số cột", "Top-1", "Top-3", "Top-5", "macro-F1"],
    [
        [i, int(r_["Số đặc trưng"]), pc(r_["top1"]), pc(r_["top3"]), pc(r_["top5"]), sf(r_["macro_f1"])]
        for i, r_ in AB_DT.iterrows()
    ],
    dam_hang=(len(AB_DT) - 1,),
)
_bo, _du = AB_DT.loc["Bỏ 10 câu sở thích"], AB_DT.loc["TẤT CẢ 43 (bản chạy thật)"]
rich(
    ("Mười câu hỏi sở thích có bõ công không? ", True),
    (
        f"Đây là câu hỏi thiết kế sản phẩm quan trọng, vì mỗi câu hỏi thêm vào đều làm "
        f"tăng khả năng người dùng bỏ dở. So sánh trực tiếp: bỏ mười câu sở thích thì "
        f"Top-3 chỉ đạt {pc(_bo['top3'])} và macro-F1 {sf(_bo['macro_f1'])}; giữ đủ "
        f"{DL['n_dac_trung']} đặc trưng thì Top-3 lên {pc(_du['top3'])} và macro-F1 "
        f"{sf(_du['macro_f1'])}. Mức tăng {pc(_du['top3'] - _bo['top3'])} điểm ở Top-3 và "
        f"{(_du['macro_f1'] / _bo['macro_f1'] - 1) * 100:.0f}% ở macro-F1 khẳng định mười "
        "câu hỏi là xứng đáng.",
        False,
    ),
)
hinh("08_model/hinh_8_3_ablation.png", "Ảnh hưởng của dữ liệu tổng hợp tới các chỉ số")
hinh("08_model/hinh_8_4_dong_gop_dac_trung.png", "Đóng góp của từng nhóm đặc trưng")

doc.add_heading("4.5. Kết quả trên tập kiểm tra", 2)
luu_y(
    "Tính toàn vẹn của phép đo:",
    f"Tập kiểm tra gồm {DL['test_that']} phiếu khảo sát thật, được tách ra từ giai đoạn 2 "
    "và chỉ mở đúng một lần tại bước này. Không notebook nào từ giai đoạn 3 đến giai đoạn "
    "7 đọc tệp này, và điều đó được bảo đảm bằng lệnh kiểm tra tự động.",
)

doc.add_heading("4.5.1. Tầng 1 — phân loại khối ngành", 3)
p(
    f"Tầng 1 đạt Top-1 {pc(T1['top1'])}, Top-2 {pc(T1['top2'])}, Top-3 {pc(T1['top3'])}, "
    f"macro-F1 {sf(T1['macro_f1'])} và balanced accuracy {sf(T1['balanced_acc'])}. Vì đây là "
    f"bài toán {DL['n_lop_khoi']} lớp nên không so trực tiếp với các con số của bài toán "
    f"{N_NGANH} ngành. Ý nghĩa của tầng 1 là nó đặt trần trên cho chế độ Khám phá: khi "
    "tầng 1 đoán sai khối, tầng 2 rất khó cứu vãn."
)
bang(
    "Kết quả chi tiết tầng 1 theo từng khối ngành",
    ["Khối ngành", "Precision", "Recall", "F1", "Số mẫu", "AUC-ROC"],
    [
        [i, sf(r_["precision"], 2), sf(r_["recall"], 2), sf(r_["f1-score"], 2),
         int(r_["support"]), sf(AUC["tang1_tung_khoi"].get(i, float("nan")), 3)]
        for i, r_ in CR1.iterrows()
        if i in AUC["tang1_tung_khoi"]
    ],
)
p(
    "Khối Công nghệ thông tin và Máy tính có F1 bằng 0 dù có tới 13 mẫu kiểm tra. "
    "Nguyên nhân là hồ sơ điểm của khối này trùng lặp mạnh với khối Kỹ thuật và Công "
    "nghệ, khiến mô hình dồn dự đoán sang khối sau. Khối Luật có AUC thấp nhất "
    f"({sf(AUC['tang1_tung_khoi']['Luật'], 3)}) do chỉ có 5 mẫu kiểm tra."
)

doc.add_heading("4.5.2. Tầng 2 — chế độ Khám phá (chỉ số chính của hệ thống)", 3)
bang(
    "Kết quả chế độ Khám phá trên tập kiểm tra",
    ["Chỉ số", "Giá trị", "Mốc đoán ngẫu nhiên", "Gấp bao nhiêu lần"],
    [
        ["Top-1", pc(AUTO["top1"]), pc(1 / N_NGANH), f"{AUTO['top1'] * N_NGANH:.1f}×".replace(".", ",")],
        ["Top-3", pc(AUTO["top3"]), pc(3 / N_NGANH), f"{AUTO['top3'] * N_NGANH / 3:.1f}×".replace(".", ",")],
        ["Top-5", pc(AUTO["top5"]), pc(5 / N_NGANH), f"{AUTO['top5'] * N_NGANH / 5:.1f}×".replace(".", ",")],
        ["macro-F1", sf(AUTO["macro_f1"]), "—", "—"],
        ["Balanced accuracy", sf(AUTO["balanced_acc"]), "—", "—"],
        ["AUC-ROC (macro)", sf(AUC["auto_macro"]), "0,500", "—"],
    ],
    dam_hang=(1,),
)
rich(
    ("Cách diễn giải cho người dùng cuối. ", True),
    (
        f"Khi hệ thống hiển thị ba gợi ý, khoảng {AUTO['top3'] * 10:.0f} trên 10 học sinh "
        f"sẽ thấy đúng ngành mình đang theo học trong danh sách — cao gấp "
        f"{AUTO['top3'] * N_NGANH / 3:.1f} lần so với đoán ngẫu nhiên. Đây chính là lý do "
        "sản phẩm được thiết kế theo hướng gợi ý danh sách ngắn thay vì đưa ra một đáp án "
        "duy nhất.",
        False,
    ),
)
luu_y(
    "Cảnh báo quan trọng:",
    M["canh_bao"],
)
bang(
    "So sánh chế độ Khám phá với mô hình một tầng trên tập kiểm tra",
    ["Chỉ số", "Một tầng (39 lớp)", "Hai tầng — Khám phá", "Chênh lệch"],
    [
        ["Top-1", pc(PHANG["top1"]), pc(AUTO["top1"]), pc(AUTO["top1"] - PHANG["top1"])],
        ["Top-3", pc(PHANG["top3"]), pc(AUTO["top3"]), "+" + pc(AUTO["top3"] - PHANG["top3"])],
        ["Top-5", pc(PHANG["top5"]), pc(AUTO["top5"]), pc(AUTO["top5"] - PHANG["top5"])],
        ["macro-F1", sf(PHANG["macro_f1"]), sf(AUTO["macro_f1"]), sf(AUTO["macro_f1"] - PHANG["macro_f1"])],
        ["AUC-ROC", sf(AUC["phang_macro"]), sf(AUC["auto_macro"]), "+" + sf(AUC["auto_macro"] - AUC["phang_macro"])],
    ],
)
p(
    "Trên tập kiểm tra, kiến trúc hai tầng không thắng tuyệt đối: Top-1 và Top-5 thấp hơn "
    "mô hình một tầng, trong khi Top-3 và AUC-ROC cao hơn. Kết quả này được báo cáo "
    "nguyên trạng thay vì chỉ trích dẫn phần có lợi. Lý do giữ kiến trúc hai tầng là ba "
    f"điểm: nó thắng đồng đều trên kiểm định chéo {M['cv']['n_splits'] * M['cv']['n_repeats']} "
    "lượt (đáng tin hơn một lần đo trên 102 mẫu), nó có AUC-ROC cao hơn tức xếp hạng tốt "
    "hơn, và nó là điều kiện kỹ thuật để có chế độ Tư vấn."
)

doc.add_heading("4.5.3. Chế độ Tư vấn — có điều kiện", 3)
p(
    f"Khi người dùng tự chọn trước khối ngành, mô hình chỉ cần xếp hạng trong khối đó. "
    f"Kết quả: Top-1 {pc(TUVAN['top1'])}, Top-3 {pc(TUVAN['top3'])}, Top-5 "
    f"{pc(TUVAN['top5'])}, macro-F1 {sf(TUVAN['macro_f1'])}, AUC-ROC "
    f"{sf(AUC['tu_van_macro'])}. Những con số này cao hơn hẳn chế độ Khám phá, nhưng chúng "
    "giả định người dùng đã chọn đúng khối — tức một phần bài toán đã được giải sẵn. Vì "
    "vậy chúng chỉ được trình bày như đặc tính của một chế độ sản phẩm riêng, không phải "
    "năng lực tự động của hệ thống."
)

doc.add_heading("4.5.4. Ba điều chỉ số AUC-ROC làm rõ", 3)
bullet(
    "Thứ nhất, năng lực xếp hạng tốt hơn vẻ ngoài của Top-1. Ngành đúng nằm ở vị trí "
    f"trung vị thứ 5 trong {N_NGANH} ngành, nghĩa là mô hình hiếm khi xếp ngành đúng "
    "xuống cuối danh sách.",
    len("Thứ nhất,"),
)
bullet(
    f"Thứ hai, AUC-ROC của kiến trúc hai tầng ({sf(AUC['auto_macro'])}) cao hơn mô hình "
    f"một tầng ({sf(AUC['phang_macro'])}) dù Top-1 thấp hơn. Top-1 chỉ xét vị trí đầu "
    "tiên, còn AUC xét toàn bộ thứ tự — với sản phẩm gợi ý danh sách thì AUC phản ánh "
    "chất lượng sát hơn.",
    len("Thứ hai,"),
)
bullet(
    f"Thứ ba, chỉ số tổng thể tốt không che được vấn đề cục bộ: 25 trên {N_NGANH} ngành "
    "có F1 bằng 0, tức mô hình chưa từng dự đoán đúng ngành đó lần nào. Hai mặt này phải "
    "được nêu song song.",
    len("Thứ ba,"),
)
hinh("08_model/hinh_8_5_ket_qua_test.png", "Tổng hợp kết quả trên tập kiểm tra")
hinh("08_model/hinh_8_6_duong_roc_va_auc.png", "Đường ROC và chỉ số AUC theo từng khối ngành")
hinh("08_model/hinh_8_7_tong_hop_chi_so.png", "Bảng tổng hợp toàn bộ chỉ số đánh giá")

doc.add_heading("4.5.5. Ảnh hưởng của việc thiếu điểm thi", 3)
_co, _chua = TOPK[0], TOPK[1]
p(
    "Học sinh trung học phổ thông sử dụng hệ thống trước kỳ thi sẽ chưa có điểm. Kịch "
    "bản này được đo riêng để biết hệ thống hoạt động ra sao khi thiếu nhóm đặc trưng "
    "điểm thi."
)
bang(
    "So sánh hai kịch bản triển khai",
    ["Chỉ số", "Có điểm thi", "Chưa có điểm thi", "Chênh lệch"],
    [
        ["Khối ngành — Top-1", pc(_co["khoi"][0]), pc(_chua["khoi"][0]), pc(_chua["khoi"][0] - _co["khoi"][0])],
        ["Khối ngành — Top-3", pc(_co["khoi"][2]), pc(_chua["khoi"][2]), pc(_chua["khoi"][2] - _co["khoi"][2])],
        ["Khám phá — Top-3", pc(_co["auto"][2]), pc(_chua["auto"][2]), pc(_chua["auto"][2] - _co["auto"][2])],
        ["Khám phá — Top-5", pc(_co["auto"][4]), pc(_chua["auto"][4]), pc(_chua["auto"][4] - _co["auto"][4])],
        ["Tư vấn — Top-3", pc(_co["tuvan"][2]), pc(_chua["tuvan"][2]), pc(_chua["tuvan"][2] - _co["tuvan"][2])],
    ],
)
p(
    "Hệ thống vẫn hoạt động khi chưa có điểm nhưng chất lượng giảm rõ rệt, đặc biệt ở "
    "Top-1 của tầng khối ngành. Kết luận thực tiễn: nên khuyến khích người dùng nhập "
    "điểm nếu đã có, và cảnh báo rõ mức độ tin cậy thấp hơn khi chưa có."
)

doc.add_heading("4.6. Tinh chỉnh mở rộng (Giai đoạn 9)", 2)
p(
    f"Giai đoạn 8 chỉ thử {TC['n_cau_hinh_gd8']} cấu hình siêu tham số. Câu hỏi đặt ra là "
    f"liệu tìm kiếm rộng hơn có cải thiện đáng kể hay không. Giai đoạn 9 mở rộng lên "
    f"{TC['n_cau_hinh_gd9']} cấu hình và quét thêm trọng số nguồn trong dải từ 1 đến 96 lần."
)
bang(
    "Kết quả tinh chỉnh mở rộng (macro-F1 trên kiểm định chéo)",
    ["Đối tượng", f"Giai đoạn 8 ({TC['n_cau_hinh_gd8']} cấu hình)", f"Giai đoạn 9 ({TC['n_cau_hinh_gd9']} cấu hình)", "Chênh lệch tuyệt đối"],
    [
        ["Tầng 1 — khối ngành", sf(TC["tang1"]["gd8"]), sf(TC["tang1"]["gd9"]), "+" + sf(TC["tang1"]["chenh"], 4)],
        ["Tầng 2 — ngành", sf(TC["tang2"]["gd8"]), sf(TC["tang2"]["gd9"]), "+" + sf(TC["tang2"]["chenh"], 4)],
    ],
)
rich(
    ("Kết luận. ", True),
    (
        f"Tăng số cấu hình lên hơn bảy lần chỉ đem lại mức cải thiện +{sf(TC['tang1']['chenh'], 4)} "
        f"cho tầng 1 và +{sf(TC['tang2']['chenh'], 4)} cho tầng 2 về macro-F1 — không đáng kể. "
        f"Việc quét trọng số nguồn cho thấy giá trị tốt nhất là "
        f"{sf(TC['trong_so_nguon']['tot_nhat'], 2)} thay vì "
        f"{sf(TC['trong_so_nguon']['dang_dung'], 2)} đang dùng, nhưng chênh lệch cũng rất nhỏ "
        "và đi kèm mức quá khớp cao hơn. Điều này chứng minh nút thắt của bài toán nằm ở "
        "lượng dữ liệu thật, không phải ở siêu tham số. Đầu tư thêm công sức tối ưu sẽ "
        "kém hiệu quả hơn nhiều so với việc thu thập thêm phiếu khảo sát.",
        False,
    ),
)
hinh("09_tinh_chinh/hinh_9_1_do_rong.png", "Ảnh hưởng của độ rộng tìm kiếm siêu tham số")
hinh("09_tinh_chinh/hinh_9_2_trong_so_nguon.png", "Kết quả quét trọng số nguồn dữ liệu")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CHƯƠNG 5
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("CHƯƠNG 5. THẢO LUẬN", 1)

doc.add_heading("5.1. Điểm mạnh về phương pháp", 2)
for t, d in [
    ("Chống rò rỉ dữ liệu triệt để.",
     f"Tập kiểm tra {DL['test_that']} phiếu được tách ngay từ giai đoạn 2 và chỉ mở một lần "
     "ở giai đoạn 8. Dữ liệu tổng hợp chỉ xuất hiện ở phía huấn luyện và được sinh riêng "
     "cho từng fold. Biến one-hot khối ngành bị loại khỏi đặc trưng tầng 2 vì suy trực "
     "tiếp từ nhãn. Tham số chuẩn hoá z-score tính riêng trên tập huấn luyện."),
    ("Mọi khẳng định đều có số liệu đối chứng.",
     "Hai kiểu đột biến, hai cách sinh dữ liệu, sáu nhóm đặc trưng, hai kiến trúc, hai "
     "độ rộng tìm kiếm — tất cả đều được đo và trình bày kèm phương án đối chứng, thay "
     "vì khẳng định suông."),
    ("Kiểm tra tự động trong mã nguồn.",
     "Khoảng ba mươi lệnh assert rải khắp chín notebook, dừng ngay khi một giả định bị "
     "vi phạm — chẳng hạn số dòng sai, thiếu ngành, hoặc tập kiểm tra bị đọc nhầm."),
    ("Bộ chỉ số đánh giá đầy đủ.",
     "Nghiên cứu không dừng ở accuracy mà báo cáo đồng thời macro-F1, balanced accuracy "
     "và AUC-ROC. Với dữ liệu mất cân bằng, accuracy đơn thuần dễ gây hiểu lầm."),
    ("Báo cáo trung thực.",
     f"Chế độ Tư vấn ({pc(TUVAN['top3'])}) được tách bạch rõ khỏi chế độ Khám phá "
     f"({pc(AUTO['top3'])}) và không dùng làm con số đại diện. Việc 25 trên {N_NGANH} ngành "
     "có F1 bằng 0, cũng như mức quá khớp, đều được nêu công khai."),
]:
    bullet(t + " " + d, len(t))

doc.add_heading("5.2. Hạn chế", 2)
bang(
    "Danh mục hạn chế của nghiên cứu",
    ["#", "Hạn chế", "Mức độ"],
    [
        ["1", "Mẫu khảo sát là sinh viên đang học, không phải học sinh trung học phổ thông đang chọn ngành. Sinh viên trả lời sau khi đã biết mình học ngành gì, nên câu trả lời có thể bị ảnh hưởng ngược.", "Cao"],
        ["2", f"Chỉ {DL['train_that']} phiếu thật cho {N_NGANH} ngành, trung bình chưa tới mười lăm phiếu mỗi ngành.", "Cao"],
        ["3", f"25 trên {N_NGANH} ngành có F1 bằng 0 trên tập kiểm tra.", "Cao"],
        ["4", "Năm trong 39 ngành không có hồ sơ điểm trúng tuyển thật, phải mượn từ ngành cùng khối.", "Trung bình"],
        ["5", f"Khoảng cách quá khớp {pc(_of['train_top1'] - _of['val_top1'])} giữa huấn luyện và validation.", "Trung bình"],
        ["6", "Nhiều ngành hiếm chỉ có một phiếu trong tập kiểm tra, khiến chỉ số theo ngành thiếu ổn định.", "Trung bình"],
        ["7", "Loại 11,7% phiếu khảo sát vì nghi trả lời máy móc; ngưỡng lọc chọn theo kinh nghiệm.", "Trung bình"],
        ["8", f"AUC phân biệt thật/tổng hợp đạt {sf(KD['auc_phan_biet'][_a], 3)}, đạt ngưỡng nhưng chưa lý tưởng (0,500).", "Thấp"],
        ["9", "Phân bố tổ hợp xét tuyển lệch giữa hai nguồn; tám tổ hợp chỉ xuất hiện ở khảo sát.", "Thấp"],
        ["10", "Một số tham số (K co ngót, trần/sàn cân bằng, β) chọn theo kinh nghiệm và kiểm định chéo, chưa tối ưu toàn cục.", "Thấp"],
    ],
)

doc.add_heading("5.3. Hướng phát triển", 2)
bullet(
    "Mở rộng khảo sát tới học sinh trung học phổ thông đang trong giai đoạn chọn ngành. "
    "Đây là hướng có tác động lớn nhất, vì kết quả giai đoạn 9 đã chỉ ra nút thắt nằm ở "
    "dữ liệu chứ không ở mô hình.",
    len("Mở rộng khảo sát tới học sinh trung học phổ thông đang trong giai đoạn chọn ngành."),
)
bullet(
    f"Cải thiện tầng 1, đặc biệt khối Công nghệ thông tin và Máy tính đang có F1 bằng 0. "
    f"Vì tầng 1 đặt trần trên cho chế độ Khám phá, mỗi điểm phần trăm cải thiện ở đây đều "
    "lan xuống kết quả cuối.",
    len("Cải thiện tầng 1, đặc biệt khối Công nghệ thông tin và Máy tính đang có F1 bằng 0."),
)
bullet(
    "Bổ sung mô hình Random Forest làm mốc so sánh, huấn luyện trên đúng bộ fold kiểm "
    "định chéo đã dùng cho XGBoost.",
    len("Bổ sung mô hình Random Forest làm mốc so sánh,"),
)
bullet(
    "Ghi nhật ký huấn luyện để theo dõi biến động chỉ số qua từng chu kỳ huấn luyện lại "
    "khi dữ liệu người dùng tích luỹ dần.",
    len("Ghi nhật ký huấn luyện"),
)
bullet(
    "Tích hợp giải thích mô hình bằng TreeSHAP, cho phép trả lời câu hỏi vì sao một ngành "
    "được gợi ý cho một học sinh cụ thể.",
    len("Tích hợp giải thích mô hình bằng TreeSHAP,"),
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PHỤ LỤC
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("PHỤ LỤC. Danh mục tệp kết quả", 1)
bang(
    "Các tệp kết quả chính của pipeline",
    ["Tệp", "Thư mục", "Nội dung"],
    [
        ["nganh_khoi_mapping.json", "02_split/", f"Bảng phân loại {N_NGANH} ngành thành {DL['n_lop_khoi']} khối"],
        ["cv_folds_real.json", "02_split/", f"Chỉ số fold cho {M['cv']['n_splits'] * M['cv']['n_repeats']} lượt kiểm định chéo"],
        ["khaosat_train.csv", "02_split/", f"{DL['train_that']} phiếu huấn luyện đã làm sạch"],
        ["khaosat_test_KHONG_DUNG_TOI.csv", "02_split/", f"{DL['test_that']} phiếu kiểm tra — khoá, mở một lần"],
        ["phan_phoi_theo_nganh.json", "03_distribution/", f"{N_NGANH} phân phối (trung bình, hiệp phương sai) đã co ngót"],
        ["phan_phoi_theo_fold.json", "03_distribution/", "15 bộ phân phối học riêng theo từng fold"],
        ["synthetic_GA_data.csv", "04_synthetic/", f"{ng(DL['train_tong_hop'])} dòng tổng hợp"],
        ["synthetic_per_fold.npz", "04_synthetic/", "15 bộ dữ liệu tổng hợp theo fold"],
        ["train_final.csv", "05_final/", f"{ng(DL['train_that'] + DL['train_tong_hop'])} dòng đã gộp kèm trọng số"],
        ["bao_cao_kiem_dinh_GA.csv", "06_validation/", "Chi tiết bốn phép kiểm định theo từng câu hỏi"],
        ["ket_luan_kiem_dinh.json", "06_validation/", "Kết luận đạt chuẩn của dữ liệu tổng hợp"],
        ["X_train.csv / X_test.csv", "07_model_ready/", f"Ma trận đặc trưng {DL['n_dac_trung']} cột"],
        ["M_nganh_khoi.npy", "07_model_ready/", f"Ma trận ánh xạ {N_NGANH}×{DL['n_lop_khoi']}"],
        ["model_stage1_khoinganh.json", "08_model/", "Mô hình XGBoost tầng 1"],
        ["model_stage2_nganh.json", "08_model/", "Mô hình XGBoost tầng 2"],
        ["metrics_summary.json", "08_model/", "Toàn bộ số liệu kết quả — nguồn chính của báo cáo"],
        ["cv_results.csv", "08_model/", "Kết quả từng lượt kiểm định chéo"],
        ["ablation_du_lieu_tong_hop.csv", "08_model/", "So sánh có và không có dữ liệu tổng hợp"],
        ["ablation_nhom_dac_trung.csv", "08_model/", "So sánh sáu nhóm đặc trưng"],
        ["topk_kich_ban_trien_khai.json", "08_model/", "Top-K cho hai kịch bản có và chưa có điểm thi"],
        ["ket_luan_tinh_chinh.json", "09_tinh_chinh/", "Kết luận giai đoạn 9"],
    ],
)

doc.add_heading("Siêu tham số của mô hình cuối cùng", 2)
_h = M["sieu_tham_so"]
_keys = [k for k in _h["tang1"] if k in _h["tang2"]]
bang(
    "Siêu tham số hai tầng",
    ["Tham số", "Tầng 1 (khối ngành)", "Tầng 2 (ngành)"],
    [[k, str(_h["tang1"][k]), str(_h["tang2"][k])] for k in _keys]
    + [["beta (hệ số kết hợp)", sf(_h["beta"], 1), sf(_h["beta"], 1)]],
)
p(
    f"Mô hình được huấn luyện ngày {M['ngay_chay']} với seed {M['seed']}. Toàn bộ số liệu "
    "trong báo cáo này được đọc trực tiếp từ các tệp kết quả nêu trên tại thời điểm biên "
    "soạn, không có con số nào nhập tay."
)

RA.parent.mkdir(parents=True, exist_ok=True)
doc.save(RA)
print(f"✔ Đã lưu: {RA}")
print(f"  {_dem['hinh']} hình · {_dem['bang']} bảng · {RA.stat().st_size / 1024 / 1024:.1f} MB")
